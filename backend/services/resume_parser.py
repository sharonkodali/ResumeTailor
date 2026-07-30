"""
Turn an uploaded resume file into plain text.

This layer is deliberately dumb: it only recovers readable text and never tries
to infer structure. Structuring is the LLM's job in resume_extractor.py.
"""

import io
import re
from pathlib import Path

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".tex")

# Resumes are short. Anything past this is a runaway file, not a resume, and
# there is no reason to pay to send it to the model.
MAX_CHARS = 20_000


class ResumeParseError(Exception):
    """Raised when an uploaded file cannot be turned into usable text."""


def extract_text(filename: str, data: bytes) -> str:
    """
    Dispatch on file extension and return normalized plain text.

    Raises ResumeParseError for unsupported types, corrupt files, and files
    that yield no text (a scanned-image PDF is the common case).
    """
    suffix = Path(filename or "").suffix.lower()

    if suffix == ".pdf":
        text = _from_pdf(data)
    elif suffix == ".docx":
        text = _from_docx(data)
    elif suffix == ".tex":
        text = _from_latex(data)
    else:
        raise ResumeParseError(
            f"Unsupported file type '{suffix or filename}'. "
            f"Upload one of: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    text = _normalize(text)

    if not text.strip():
        hint = (
            " If it is a scanned or image-based PDF, export a text-based version."
            if suffix == ".pdf"
            else ""
        )
        raise ResumeParseError(f"No text could be read from that file.{hint}")

    return text[:MAX_CHARS]


# ---------------------------------------------------------------------------
# Per-format extraction
# ---------------------------------------------------------------------------

def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    from pypdf.errors import PyPdfError

    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            # An empty user password unlocks most "protected" resumes.
            try:
                reader.decrypt("")
            except Exception as exc:
                raise ResumeParseError(
                    "That PDF is password protected. Remove the password and try again."
                ) from exc
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except ResumeParseError:
        raise
    except (PyPdfError, OSError, ValueError) as exc:
        raise ResumeParseError(f"Could not read that PDF: {exc}") from exc


def _from_docx(data: bytes) -> str:
    import zipfile

    import docx
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = docx.Document(io.BytesIO(data))
    except (PackageNotFoundError, zipfile.BadZipFile, KeyError, ValueError) as exc:
        # A .doc renamed to .docx is the usual cause: it is not a zip container.
        raise ResumeParseError(
            "Could not read that .docx file. If it was saved as legacy .doc, "
            "re-save it as .docx and try again."
        ) from exc

    parts = [p.text for p in document.paragraphs]

    # Many resume templates lay content out in tables, which are not part of
    # document.paragraphs and would otherwise be dropped entirely.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" ".join(c for c in cells if c))

    return "\n".join(parts)


def _from_latex(data: bytes) -> str:
    """
    Strip LaTeX markup down to its text content.

    Custom resume macros (\\resumeSubheading{...}{...}) are unwrapped generically
    rather than being special-cased, so unfamiliar templates degrade to their
    argument text instead of vanishing.
    """
    src = _decode(data)

    # Work on the document body when there is a preamble to discard.
    if r"\begin{document}" in src:
        src = src.split(r"\begin{document}", 1)[1]
    src = src.split(r"\end{document}", 1)[0]

    # Comments: a % that is not escaped runs to end of line.
    src = re.sub(r"(?<!\\)%.*", "", src)

    # Declaration-style commands carry no resume content.
    src = re.sub(
        r"\\(?:usepackage|documentclass|newcommand|renewcommand|providecommand"
        r"|definecolor|setlength|titleformat|titlespacing|pagestyle|input"
        r"|RequirePackage|hypersetup|geometry|label|vspace|hspace|raisebox)\b"
        r"\s*(?:\[[^\]]*\])?(?:\{[^{}]*\})*",
        "",
        src,
    )

    # \href{url}{label} should keep the label, not the URL.
    src = re.sub(r"\\href\s*\{[^{}]*\}\s*\{([^{}]*)\}", r"\1", src)

    src = re.sub(r"\\item\b", "\n- ", src)          # bullets
    src = re.sub(r"\\\\(?:\[[^\]]*\])?", "\n", src)  # forced line breaks
    src = re.sub(r"\\begin\s*\{[^{}]*\}(?:\[[^\]]*\])?", "\n", src)
    src = re.sub(r"\\end\s*\{[^{}]*\}", "\n", src)
    src = re.sub(r"(?<!\\)&", " ", src)              # tabular column separator

    # Unwrap \cmd{body} -> body. The regex only matches commands whose argument
    # contains no braces, so repeating it peels nested calls from the inside out.
    for _ in range(10):
        unwrapped = re.sub(
            r"\\[a-zA-Z@]+\*?\s*(?:\[[^\]]*\])?\s*\{([^{}]*)\}", r"\1", src
        )
        if unwrapped == src:
            break
        src = unwrapped

    src = re.sub(r"\\[a-zA-Z@]+\*?", "", src)        # argument-less commands
    src = re.sub(r"\\([&%$#_{}])", r"\1", src)       # escaped literals
    src = src.replace("~", " ").replace("{", " ").replace("}", " ")

    # Typographic leftovers that would otherwise reach the model as noise.
    src = src.replace("---", "-").replace("--", "-")
    src = re.sub(r"\$([^$]*)\$", r"\1", src)         # inline math delimiters

    return src


def _decode(data: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _normalize(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    # Drop bullet markers left empty once their macro was stripped.
    text = re.sub(r"^-\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
